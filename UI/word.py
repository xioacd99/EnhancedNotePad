import os,sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docxcompose.composer import Composer

import qtawesome
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtGui import QCursor
from PyQt5.QAxContainer import QAxWidget
from PyQt5.QtWidgets import QMainWindow, QGridLayout, QWidget, QPushButton, QFileDialog, QTextEdit, QMessageBox


class Doc:
    def __init__(self,file_name=None):
        if os.path.exists(file_name):
            self.file_name = file_name
            self.documen
            t=Document(file_name)
        else:
            self.file_name='python_code.docx'
            self.document=Document()

    def ClearText(self):
        # doc=Document()
        # doc.save(self.file_name)
        pass

    #添加标题
    def AddHeadText(self,text, size):
        title_ = self.document.add_heading(level=3)
        title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
        title_run = title_.add_run(text)  # 添加标题内容
        title_run.font.size = Pt(size)  # 设置标题字体大小
        title_run.font.name = 'Times New Roman'  # 设置标题西文字体
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置标题中文字体
        title_run.font.color.rgb = RGBColor(0, 0, 0)#字体颜色

    #添加段落内容(参数1：文本内容，参数2：字体大小，参数3：上行距,参数4：字体粗细，参数5：段落位置)
    def AddParaText(self,text, size, space, thickness, position,front_name):
        p = self.document.add_paragraph()  # 段落
        p.paragraph_format.line_spacing = Pt(20)  # 行距20磅
        #判断居中还是靠左,0为靠左
        if position == 0:
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT #靠左
        else :
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
        p.paragraph_format.space_before = Pt(space)
        text = p.add_run(text)
        #判断字体是否加粗（1为不加粗）
        if thickness == 1:
            text.bold = False
        else:
            text.bold = True #加粗
        text.font.name = 'Times New Roman'
        text.element.rPr.rFonts.set(qn('w:eastAsia'), front_name)
        text.font.size = Pt(size)
        self.document.save(self.file_name)
        return p

    def AddRunText(self,p,text,bond=False,size=12):#默认12磅小四号
        r=p.add_run(text)
        r.font.name='Times New Roman'
        r.font.size = Pt(size)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.document.save(self.file_name)


class WordUi(QMainWindow):

    def __init__(self):
        super().__init__()
        self.initUI()

        self.headline_1={}
        self.headline_2 = {}
        self.headline_3 = {}
        self.text = {}
        self.headline=[]

        self.count=0
        self.headline_1_cnt=0
        self.headline_2_cnt = 0
        self.headline_3_cnt = 0
        self.text_cnt = 0

    def initUI(self):
        self.setFixedSize(960, 700)
        self.main_widget = QWidget()  # 创建窗口主部件
        self.main_layout = QGridLayout()  # 创建主部件的网格布局
        self.main_widget.setLayout(self.main_layout)  # 设置窗口主部件布局为网格布局

        self.left_widget = QWidget()  # 创建左侧部件
        self.left_widget.setObjectName('left_widget')
        self.left_layout = QGridLayout()  # 创建左侧部件的网格布局
        self.left_widget.setLayout(self.left_layout)  # 设置左侧部件布局为网格布局

        self.right_widget = QWidget()  # 创建右侧部件
        self.right_widget.setObjectName('right_widget')
        self.right_layout = QGridLayout()  # 创建右侧部件的网格布局
        self.right_widget.setLayout(self.right_layout)  # 设置右侧部件布局为网格布局

        self.main_layout.addWidget(self.left_widget, 0, 0, 12, 2)
        self.main_layout.addWidget(self.right_widget, 0, 2, 12, 10)
        self.setCentralWidget(self.main_widget)  # 设置窗口主部件



        self.button1 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加题目")  # 创建按钮1
        self.left_layout.addWidget(self.button1, 2, 0, 1, 3)
        self.button1.clicked.connect(self.add_title)
        self.button2 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加摘要、关键词")  # 创建按钮2
        self.left_layout.addWidget(self.button2, 4, 0, 1, 3)
        self.button2.clicked.connect(self.add_keywords)
        self.button3 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加1级标题")  # 创建按钮3
        self.left_layout.addWidget(self.button3, 6, 0, 1, 3)
        self.button3.clicked.connect(self.add_headline_1)
        self.button4 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加2级标题")  # 创建按钮4
        self.left_layout.addWidget(self.button4, 8, 0, 1, 3)
        self.button4.clicked.connect(self.add_headline_2)
        self.button5 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加3级标题")  # 创建按钮5
        self.left_layout.addWidget(self.button5, 10, 0, 1, 3)
        self.button5.clicked.connect(self.add_headline_3)
        self.button6 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加正文")  # 创建按钮6
        self.left_layout.addWidget(self.button6, 12, 0, 1, 3)
        self.button6.clicked.connect(self.add_text)
        self.button7 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "添加参考文献")  # 创建按钮7
        self.left_layout.addWidget(self.button7, 14, 0, 1, 3)
        self.button7.clicked.connect(self.add_reference)
        self.button8 = QPushButton(qtawesome.icon('fa.sellsy', color='white'), "生成并添加封面")  # 创建按钮8
        self.left_layout.addWidget(self.button8, 16, 0, 1, 3)
        self.button8.clicked.connect(self.generate)

        #美化
        self.left_close = QtWidgets.QPushButton("")  # 关闭按钮
        self.left_visit = QtWidgets.QPushButton("")  # 空白按钮
        self.left_mini = QtWidgets.QPushButton("")  # 最小化按钮
        #按钮绑定相应事件
        self.left_close.clicked.connect(self.close)
        self.left_visit.clicked.connect(self.showMaximized)
        self.left_mini.clicked.connect(self.showMinimized)
        #分配位置
        self.left_layout.addWidget(self.left_mini, 0, 0, 1, 1)
        self.left_layout.addWidget(self.left_visit, 0, 1, 1, 1)
        self.left_layout.addWidget(self.left_close, 0, 2, 1, 1)
        #QSS
        self.left_widget.setStyleSheet('''
          QPushButton{border:none;color:white;}
          QPushButton#left_label{
            border:none;
            border-bottom:1px solid white;
            font-size:18px;
            font-weight:700;
            font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
          }
          QPushButton#left_button:hover{border-left:4px solid red;font-weight:700;}
        ''')
        self.left_close.setStyleSheet(
            '''QPushButton{background:#F76677;border-radius:5px;}QPushButton:hover{background:red;}''')
        self.left_visit.setStyleSheet(
            '''QPushButton{background:#F7D674;border-radius:5px;}QPushButton:hover{background:yellow;}''')
        self.left_mini.setStyleSheet(
            '''QPushButton{background:#6DDF6D;border-radius:5px;}QPushButton:hover{background:green;}''')
        self.left_close.setFixedSize(15,15) # 设置关闭按钮的大小
        self.left_visit.setFixedSize(15, 15) # 设置按钮大小
        self.left_mini.setFixedSize(15, 15) # 设置最小化按钮大小
        self.main_widget.setStyleSheet('''
        QWidget#left_widget{
        background:gray;
        border-top:1px solid white;
        border-bottom:1px solid white;
        border-left:1px solid white;
        border-top-left-radius:10px;
        border-bottom-left-radius:10px;
        }
        ''')
        self.setWindowOpacity(0.9)  # 设置窗口透明度
        #self.setAttribute(QtCore.Qt.WA_TranslucentBackground)  # 设置窗口背景透明
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)  # 隐藏边框
        self.main_layout.setSpacing(0) #空隙
    #3个函数实现窗口移动
    def mousePressEvent(self, event):
        try:
            if event.button() == QtCore.Qt.LeftButton:
                self.m_flag = True
                self.m_Position = event.globalPos() - self.pos()  # 获取鼠标相对窗口的位置
                event.accept()
                self.setCursor(QCursor(QtCore.Qt.OpenHandCursor))  # 更改鼠标图标
        except Exception as e:
            print(e)
    def mouseMoveEvent(self, QMouseEvent):
        try:
            if QtCore.Qt.LeftButton and self.m_flag:
                self.move(QMouseEvent.globalPos() - self.m_Position)  # 更改窗口位置
                QMouseEvent.accept()
        except Exception as e:
            print(e)
    def mouseReleaseEvent(self, QMouseEvent):
        try:
            self.m_flag = False
            self.setCursor(QCursor(QtCore.Qt.ArrowCursor))
        except Exception as e:
            print(e)

    def add_title(self):
        self.title = QtWidgets.QLineEdit()
        self.title.setPlaceholderText("请输入题目")
        self.right_layout.addWidget(self.title,self.count,0,1,8)
        self.count+=1
    def add_keywords(self):
        self.abstract = QtWidgets.QLineEdit()
        self.abstract.setPlaceholderText("请输入摘要")
        self.right_layout.addWidget(self.abstract, self.count, 0, 1, 8)
        self.keywords = QtWidgets.QLineEdit()
        self.keywords.setPlaceholderText("请输入关键词")
        self.right_layout.addWidget(self.keywords,self.count+1,0,1,8)
        self.count += 2
    def add_headline_1(self):
        try:
            self.headline_1[self.headline_1_cnt] = QtWidgets.QLineEdit()
            self.headline_1[self.headline_1_cnt].setPlaceholderText("请输入1级标题")
            self.right_layout.addWidget(self.headline_1[self.headline_1_cnt],self.count,0,1,8)
            self.headline.append(self.headline_1[self.headline_1_cnt])
            self.count += 1
            self.headline_1_cnt+=1
        except Exception as e:
            print(e)
    def add_headline_2(self):
        try:
            self.headline_2[self.headline_2_cnt] = QtWidgets.QLineEdit()
            self.headline_2[self.headline_2_cnt].setPlaceholderText("请输入2级标题")
            self.right_layout.addWidget(self.headline_2[self.headline_2_cnt],self.count,1,1,8)
            self.headline.append(self.headline_2[self.headline_2_cnt])
            self.count += 1
            self.headline_2_cnt += 1
        except Exception as e:
            print(e)
    def add_headline_3(self):
        self.headline_3[self.headline_3_cnt] = QtWidgets.QLineEdit()
        self.headline_3[self.headline_3_cnt].setPlaceholderText("请输入3级标题")
        self.right_layout.addWidget(self.headline_3[self.headline_3_cnt], self.count, 2, 1, 8)
        self.headline.append(self.headline_3[self.headline_3_cnt])
        self.count += 1
        self.headline_3_cnt += 1
    def add_text(self):
        self.text[self.text_cnt] = QtWidgets.QLineEdit()
        self.text[self.text_cnt].setPlaceholderText("请输入正文")
        self.right_layout.addWidget(self.text[self.text_cnt], self.count, 3, 1, 8)
        self.headline.append(self.text[self.text_cnt])
        self.count += 1
        self.text_cnt += 1
    def add_reference(self):
        try:
            self.reference = QtWidgets.QLineEdit()
            self.reference.setPlaceholderText("请输入参考文献")
            self.right_layout.addWidget(self.reference,self.count,0,1,8)
            self.count += 1
        except Exception as e:
            print(e)
    def generate(self):
        try:
            #获取内容
            if self.title:
                title_text=self.title.text()
            if self.keywords and self.abstract:
                keywords_text=self.keywords.text()
                abstract_text=self.abstract.text()
            if self.reference:
                reference_text=self.reference.text()

            # headline=[(self.category(h),h.text()) for h in self.headline]
            # print(title_text,keywords_text,reference_text)
            # print(headline)

            #填写word
            file_name='python_word.docx'
            doc=Doc(file_name)
            doc.AddParaText(title_text,22,18,0,1,'隶书')#标题
            zy=doc.AddParaText('摘 要：',12,18,0,0,'宋体')
            doc.AddRunText(zy,abstract_text)
            kw=doc.AddParaText('关键字：',12,18,0,0,'宋体')
            doc.AddRunText(kw,keywords_text)

            #正文
            for h in self.headline:
                if self.category(h) == 1:
                    doc.AddParaText(f'{h.text()}',16,18,0,0,'宋体')
                elif self.category(h) == 2:
                    doc.AddParaText(f'{h.text()}',14,18,0,0,'宋体')
                elif self.category(h) == 3:
                    doc.AddParaText(f'{h.text()}',12,18,0,0,'宋体')
                elif self.category(h) == 4:
                    doc.AddParaText(f'{h.text()}',12,18,1,0,'宋体')

            re=doc.AddParaText('参考文献\n',16,18,0,0,'宋体')
            doc.AddRunText(re, reference_text,10.5)

            #添加封面
            combine_word_documents('cover.docx','python_code.docx')
            # self.openOffice("combined.docx", 'Word.Application')

        except Exception as e:
            print(e)
    def category(self,headline):
        if headline in self.headline_1.values():
            return 1
        elif headline in self.headline_2.values():
            return 2
        elif headline in self.headline_3.values():
            return 3
        else:
            return 4

    def openOffice(self, path, app):
        self.axWidget = QAxWidget(self)
        self.right_layout.addWidget(self.axWidget)
        self.axWidget.clear()
        if not self.axWidget.setControl(app):
            return QMessageBox.critical(self, '错误', '没有安装  %s' % app)
        self.axWidget.dynamicCall(
            'SetVisible (bool Visible)', 'false')  # 不显示窗体
        self.axWidget.setProperty('DisplayAlerts', False)
        self.axWidget.setControl(path)


def combine_word_documents(file_name1,file_name2):
    file1 = Document(file_name1)
    file2 = Document(file_name2)
    composer = Composer(file1)
    composer.append(file2)
    composer.save("combined.docx")



if __name__=='__main__':
    # file_name='python_word.docx'
    # doc=Doc(file_name)
    #
    # doc.AddParaText('标题',22,18,0,1,'隶书')#1为不加粗
    # zy=doc.AddParaText('摘 要：',12,18,0,0,'宋体')
    # doc.AddRunText(zy,'摘要里的内容')
    # kw=doc.AddParaText('关键字：',12,18,0,0,'宋体')
    # doc.AddRunText(kw,'关键字里面的内容')
    #
    #
    # yi=doc.AddParaText('一级标题\n',16,18,0,0,'宋体')
    # doc.AddRunText(yi,'正文'*30)
    #
    #
    # doc.AddParaText('参考文献',10.5,18,0,0,'宋体')
    # doc.ClearText()
    app = QtWidgets.QApplication(sys.argv)
    gui = WordUi()
    gui.show()

    sys.exit(app.exec_())
